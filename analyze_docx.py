from docx import Document
import sys
doc = Document(sys.argv[1])
print(f"=== {sys.argv[1].split('/')[-1]} ===")
print(f"\n[段落样式] 共{len(doc.paragraphs)}段，前30段:")
for i, p in enumerate(doc.paragraphs[:30]):
    if p.text.strip() or i < 15:
        print(f"  {i}: [{p.style.name}] {p.text[:55]}")
print(f"\n[表格] 共{len(doc.tables)}个:")
for ti, t in enumerate(doc.tables):
    print(f"  表{ti}: {len(t.rows)}行x{len(t.columns)}列 -> 第一行: {[c.text[:12] for c in t.rows[0].cells]}")
print(f"\n[节/页眉页脚] 共{len(doc.sections)}节:")
for si, s in enumerate(doc.sections):
    hdr = [p.text for p in s.header.paragraphs if p.text.strip()]
    ftr = [p.text for p in s.footer.paragraphs if p.text.strip()]
    print(f"  节{si}: 页眉={hdr} 页脚={ftr}")
print(f"\n[样式定义] 文档内置样式:")
for st in doc.styles:
    if 'Head' in st.name or 'Normal' in st.name or 'TOC' in st.name or '标题' in st.name:
        print(f"  {st.name} ({st.type})")
