import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Twips
import build_docx


def _tpl_with_custom_styles(path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(11906); sec.left_margin = Twips(1440); sec.right_margin = Twips(1440)
    for nm in ('报告标题', '编制日期'):
        doc.styles.add_style(nm, WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph('旧标题', style='报告标题')
    # 日期占位含 '年' 以满足 build() 的日期替换触发条件
    doc.add_paragraph('2020年1月', style='编制日期')
    t = doc.add_table(rows=2, cols=5)
    t.rows[0].cells[0].text = '版本号'
    doc.add_paragraph('旧正文', style='Heading 1')
    doc.save(str(path))


def test_build_uses_manifest_styles(tmp_path):
    tpl = tmp_path / 'tpl.docx'; _tpl_with_custom_styles(str(tpl))
    md = tmp_path / 'a.md'; md.write_text('# 概述\n正文段\n', encoding='utf-8')
    out = tmp_path / 'o.docx'
    manifest = {
        'cover_title_style': '报告标题',
        'cover_date_style': '编制日期',
        'heading_styles': ['Heading 1', 'Heading 2', 'Heading 3'],
        'caption_styles': {'table': None, 'figure': None},
        'change_log_table': 'auto',
        'strip_heading_num': False,
    }
    build_docx.build(str(tpl), str(md), str(out),
                     cover_title='新标题', cover_date='2026年1月',
                     change_log_rows=[['V1.0', '2026.01', '组', '增加', '初建']],
                     manifest=manifest)
    d = Document(str(out))
    texts = [p.text for p in d.paragraphs]
    assert '新标题' in texts
    assert '2026年1月' in texts
    assert d.tables[0].rows[1].cells[0].text == 'V1.0'


def test_build_manifest_explicit_null_disables_date(tmp_path):
    # manifest 显式 cover_date_style: null → 不填日期（resolve_style 返回 None）
    tpl = tmp_path / 'tpl.docx'; _tpl_with_custom_styles(str(tpl))
    md = tmp_path / 'a.md'; md.write_text('# 概述\n正文\n', encoding='utf-8')
    out = tmp_path / 'o.docx'
    manifest = {
        'cover_title_style': '报告标题',
        'cover_date_style': None,   # 显式禁用
        'heading_styles': ['Heading 1', 'Heading 2', 'Heading 3'],
        'change_log_table': 'auto',
        'strip_heading_num': False,
    }
    build_docx.build(str(tpl), str(md), str(out),
                     cover_title='新标题', cover_date='2026年1月', manifest=manifest)
    d = Document(str(out))
    texts = [p.text for p in d.paragraphs]
    assert '新标题' in texts
    assert '2026年1月' not in texts   # 被禁用，日期没填进去


def test_build_manifest_absent_falls_back(tmp_path):
    # manifest 不含 cover_title_style → find_style 兜底找候选（模板有'报告标题'但不在候选里→None→warn跳过）
    tpl = tmp_path / 'tpl.docx'; _tpl_with_custom_styles(str(tpl))
    md = tmp_path / 'a.md'; md.write_text('# 概述\n正文\n', encoding='utf-8')
    out = tmp_path / 'o.docx'
    build_docx.build(str(tpl), str(md), str(out),
                     cover_title='新标题', cover_date='2026年1月', manifest={})
    d = Document(str(out))
    # 候选里没有'报告标题'，兜底找不到→跳过封面标题填入（不报错即可）
    # 正文应正常生成
    assert any(p.text == '概述' for p in d.paragraphs) or any('概述' in p.text for p in d.paragraphs)


def _tpl_with_chinese_headings(path):
    """模板用中文样式名 标题 1/2/3，且 标题 1 带 numPr（自动编号）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(11906); sec.left_margin = Twips(1440); sec.right_margin = Twips(1440)
    for nm in ('报告标题', '编制日期'):
        doc.styles.add_style(nm, WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph('旧标题', style='报告标题')
    doc.add_paragraph('2020年1月', style='编制日期')
    t = doc.add_table(rows=2, cols=5)
    t.rows[0].cells[0].text = '版本号'
    # 中文标题样式；给 标题 1 加 numPr 模拟自动编号
    h1 = doc.styles.add_style('标题 1', WD_STYLE_TYPE.PARAGRAPH)
    pPr = h1.element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0'); numPr.append(ilvl)
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1'); numPr.append(numId)
    pPr.append(numPr)
    doc.styles.add_style('标题 2', WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style('标题 3', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph('旧正文', style='标题 1')
    doc.save(str(path))


def test_build_chinese_heading_styles_and_strip(tmp_path):
    """中文样式名 标题 1/2/3 能用，且 strip_heading_num=True 能去掉中文标题样式的 numPr。"""
    tpl = tmp_path / 'tpl.docx'; _tpl_with_chinese_headings(str(tpl))
    md = tmp_path / 'a.md'; md.write_text('# 概述\n正文\n', encoding='utf-8')
    out = tmp_path / 'o.docx'
    manifest = {
        'cover_title_style': '报告标题',
        'cover_date_style': '编制日期',
        'heading_styles': ['标题 1', '标题 2', '标题 3'],
        'caption_styles': {'table': None, 'figure': None},
        'change_log_table': 'auto',
        'strip_heading_num': True,
    }
    build_docx.build(str(tpl), str(md), str(out),
                     cover_title='新标题', cover_date='2026年1月', manifest=manifest)
    d = Document(str(out))
    # h1 用了中文样式 标题 1
    h1_paras = [p for p in d.paragraphs if p.text.strip() == '概述']
    assert h1_paras, '概述 段未生成'
    assert h1_paras[0].style.name == '标题 1'
    # strip 生效：标题 1 样式的 numPr 被移除
    from docx.oxml.ns import qn
    s = d.styles['标题 1']
    pPr = s.element.find(qn('w:pPr'))
    assert pPr is None or pPr.find(qn('w:numPr')) is None, 'strip 后 标题 1 仍含 numPr'
