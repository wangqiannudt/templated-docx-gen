from docx import Document
from docx.shared import Twips
import build_docx


def test_detect_body_width_a4_default():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(11906)
    sec.left_margin = Twips(1440)
    sec.right_margin = Twips(1440)
    assert build_docx.detect_page_body_width(doc) == 11906 - 1440 - 1440


def test_detect_body_width_custom():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(10000)
    sec.left_margin = Twips(800)
    sec.right_margin = Twips(800)
    assert build_docx.detect_page_body_width(doc) == 8400


def test_detect_body_width_fallback_when_missing():
    doc = Document()
    w = build_docx.detect_page_body_width(doc)
    assert isinstance(w, int) and w > 0


def test_find_style_by_candidate():
    doc = Document()
    from docx.enum.style import WD_STYLE_TYPE
    doc.styles.add_style('报告标题', WD_STYLE_TYPE.PARAGRAPH)
    found = build_docx.find_style(doc, ['封皮标题', '报告标题', '主标题'])
    assert found == '报告标题'


def test_find_style_none_when_absent():
    doc = Document()
    assert build_docx.find_style(doc, ['不存在的样式']) is None


def test_find_style_heading_default():
    doc = Document()
    assert build_docx.find_style(doc, ['Heading 1']) == 'Heading 1'


def test_detect_change_log_table_by_header():
    doc = Document()
    t = doc.add_table(rows=2, cols=5)
    t.rows[0].cells[0].text = '版本号'
    idx, ncols = build_docx.detect_change_log_table(doc)
    assert idx == 0 and ncols == 5


def test_detect_change_log_table_keyword_variant():
    doc = Document()
    t = doc.add_table(rows=1, cols=4)
    t.rows[0].cells[0].text = '版本'
    idx, ncols = build_docx.detect_change_log_table(doc)
    assert idx == 0 and ncols == 4


def test_detect_change_log_table_none_when_no_match():
    doc = Document()
    doc.add_table(rows=2, cols=3)
    assert build_docx.detect_change_log_table(doc) is None


def test_detect_change_log_table_skips_first_picks_right():
    doc = Document()
    doc.add_table(rows=1, cols=3)
    t = doc.add_table(rows=2, cols=5)
    t.rows[0].cells[0].text = '版本号'
    idx, ncols = build_docx.detect_change_log_table(doc)
    assert idx == 1 and ncols == 5
