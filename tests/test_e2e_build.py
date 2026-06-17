import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Twips
from docx.oxml.ns import qn
import build_docx


def _make_png(path):
    from PIL import Image
    Image.new('RGB', (100, 50), 'white').save(str(path))


def _tpl(path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(11906); sec.left_margin = Twips(1440); sec.right_margin = Twips(1440)
    for nm in ('封皮标题', '封皮单位', '表题注', '图题注'):
        try:
            doc.styles.add_style(nm, WD_STYLE_TYPE.PARAGRAPH)
        except Exception:
            pass
    doc.add_paragraph('旧', style='封皮标题')
    doc.add_paragraph('2020年1月', style='封皮单位')
    t = doc.add_table(rows=2, cols=5); t.rows[0].cells[0].text = '版本号'
    doc.add_paragraph('旧', style='Heading 1')
    doc.save(str(path))


def test_e2e_full_build(tmp_path):
    # 准备 md 引用的图片（放 tmp_path/assets，md 里相对路径引用）
    assets = tmp_path / 'assets'; assets.mkdir()
    _make_png(assets / 'sample.png')
    tpl = tmp_path / 'tpl.docx'; _tpl(str(tpl))
    md = tmp_path / 'a.md'
    md.write_text(
        '# 概述\n'
        '正文段。\n'
        '图1 示意图。\n'
        '![图2 原理图](assets/sample.png)\n'
        '| 序号 | 名称 | 数量 |\n|---|---|---|\n| 1 | foo | 2 |\n',
        encoding='utf-8')
    out = tmp_path / 'o.docx'

    build_docx.build(str(tpl), str(md), str(out),
                     cover_title='测试标题', cover_date='2026年6月',
                     change_log_rows=[['V1.0', '2026.06', '组', '增加', '初建']])

    d = Document(str(out))
    xml = d.element.xml
    # SEQ 域 + dirty
    assert 'SEQ' in xml and 'fldChar' in xml
    assert 'dirty="true"' in xml or 'dirty="1"' in xml
    # 表格 tblGrid 列宽和 ≈ 版心宽(11906-1440-1440=9026)
    tbl = d.tables[0]
    grid = tbl._element.find(qn('w:tblGrid'))
    widths = [int(gc.get(qn('w:w'))) for gc in grid.findall(qn('w:gridCol'))]
    assert abs(sum(widths) - 9026) <= 5
    # 封面已填
    texts = [p.text for p in d.paragraphs]
    assert '测试标题' in texts and '2026年6月' in texts
    # 变更记录 5 列重建
    assert d.tables[0].rows[1].cells[0].text == 'V1.0'
    # 图片已嵌入（drawing 元素存在）
    assert 'graphic' in xml or 'pic:pic' in xml or '<w:drawing' in xml
