import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Twips
import analyze_template


def test_analyze_minimal_template(tmp_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(11906); sec.left_margin = Twips(1440); sec.right_margin = Twips(1440)
    st = doc.styles.add_style('封皮标题', WD_STYLE_TYPE.PARAGRAPH)
    st.font.size = Pt(22)
    doc.add_paragraph('占位标题', style='封皮标题')
    doc.styles.add_style('封皮单位', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph('2025年12月', style='封皮单位')
    t = doc.add_table(rows=2, cols=5)
    t.rows[0].cells[0].text = '版本号'
    doc.add_paragraph('正文', style='Heading 1')
    path = tmp_path / 'tpl.docx'
    doc.save(str(path))

    m = analyze_template.analyze(str(path))
    assert m['cover_title_style'] == '封皮标题'
    assert m['cover_date_style'] == '封皮单位'
    assert m['heading_styles'][0] == 'Heading 1'
    assert m['change_log_table'] == 'auto'
    assert m['strip_heading_num'] is False


def test_analyze_emits_yaml(tmp_path):
    doc = Document()
    doc.add_paragraph('正文', style='Heading 1')
    path = tmp_path / 'tpl.docx'
    doc.save(str(path))
    text = analyze_template.analyze_yaml(str(path))
    data = yaml.safe_load(text)
    assert 'heading_styles' in data
    assert data['change_log_table'] == 'auto'


def test_analyze_custom_title_style_by_fontsize(tmp_path):
    # 封面标题样式名非默认候选，靠字号最大命中
    doc = Document()
    st = doc.styles.add_style('报告主标题', WD_STYLE_TYPE.PARAGRAPH)
    st.font.size = Pt(26)
    doc.add_paragraph('大标题', style='报告主标题')
    doc.add_paragraph('正文', style='Heading 1')
    path = tmp_path / 'tpl.docx'
    doc.save(str(path))
    m = analyze_template.analyze(str(path))
    assert m['cover_title_style'] == '报告主标题'


def test_analyze_dashdash_o_writes_file(tmp_path):
    import sys
    doc = Document()
    doc.add_paragraph('正文', style='Heading 1')
    tpl = tmp_path / 'tpl.docx'
    doc.save(str(tpl))
    out = tmp_path / 'm.yaml'
    import subprocess, os
    venv_py = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'docenv', 'bin', 'python')
    r = subprocess.run([venv_py, 'analyze_template.py', str(tpl), '-o', str(out)],
                       cwd=os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                       capture_output=True, text=True)
    assert r.returncode == 0, r.stderr
    assert out.exists()
    data = yaml.safe_load(out.read_text(encoding='utf-8'))
    assert 'heading_styles' in data
