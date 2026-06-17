#!/usr/bin/env python3
"""探测模板 docx 结构，生成 manifest 草稿 YAML（供用户确认后存盘复用）。"""
import re
import argparse
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def _first_h1_index(doc):
    for i, p in enumerate(doc.paragraphs):
        try:
            if p.style.name in ('Heading 1', '标题 1'):
                return i
        except Exception:
            pass
    return len(doc.paragraphs)


def _existing_style_names(doc):
    """只收段落样式名（排除字符样式）。
    Word 为每个段落样式生成关联字符样式 '<名> 字符'，若一并收集会令 substring
    匹配（如 find_caption）误命中字符样式（'表题注 字符'），故只取段落样式。"""
    names = set()
    for s in doc.styles:
        try:
            if s.type == WD_STYLE_TYPE.PARAGRAPH:
                names.add(s.name)
        except Exception:
            pass
    return names


def analyze(template_path):
    """探测模板，返回 manifest dict。"""
    doc = Document(template_path)
    paras = doc.paragraphs
    h1i = _first_h1_index(doc)
    front = paras[:h1i]
    existing = _existing_style_names(doc)

    def font_size(p):
        try:
            return p.style.font.size.pt if p.style.font.size else 0
        except Exception:
            return 0

    # 封面主标题：front 区非空、字号最大的段样式（兜底候选名）
    title_p = None
    best = 0
    for p in front:
        if not p.text.strip():
            continue
        sz = font_size(p)
        if sz > best:
            best = sz
            title_p = p
    cover_title_style = None
    if title_p:
        try:
            cover_title_style = title_p.style.name
        except Exception:
            cover_title_style = None
    if not cover_title_style:
        for c in ('封皮标题', '报告标题', '文档标题', '主标题', '标题'):
            if c in existing:
                cover_title_style = c
                break

    # 封面日期：front 区文本含 年/年份/四位数字 的段样式
    cover_date_style = None
    for p in front:
        if p.text.strip() and ('年' in p.text or re.search(r'\d{4}', p.text)):
            try:
                cover_date_style = p.style.name
                break
            except Exception:
                pass

    # 正文 Heading 样式
    def find_h(level):
        for c in (f'Heading {level}', f'标题 {level}'):
            if c in existing:
                return c
        return f'Heading {level}'
    heading_styles = [find_h(1), find_h(2), find_h(3)]

    # 题注样式
    def find_caption(keywords):
        for n in existing:
            if any(k in n for k in keywords):
                return n
        return None
    caption_table = find_caption(['表题注', '表题', '表格题注'])
    caption_figure = find_caption(['图题注', '图题', '图片题注'])

    return {
        'cover_title_style': cover_title_style,
        'cover_date_style': cover_date_style,
        'heading_styles': heading_styles,
        'caption_styles': {'table': caption_table, 'figure': caption_figure},
        'change_log_table': 'auto',
        'strip_heading_num': False,
    }


def analyze_yaml(template_path):
    import yaml
    return yaml.safe_dump(analyze(template_path), allow_unicode=True, sort_keys=False)


def main():
    p = argparse.ArgumentParser(description='探测模板结构生成 manifest 草稿')
    p.add_argument('template')
    p.add_argument('-o', '--output', default=None, help='写入文件；省略则打印 stdout')
    a = p.parse_args()
    text = analyze_yaml(a.template)
    if a.output:
        with open(a.output, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f'已写入 manifest 草稿: {a.output}')
    else:
        print(text)


if __name__ == '__main__':
    main()
