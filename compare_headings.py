from docx import Document
import sys
for label, path in [('【v0.0模板】', sys.argv[1]), ('【我生成的】', sys.argv[2])]:
    doc = Document(path)
    print(f'\n{label} {path.split("/")[-1]}')
    print('  标题段落（Heading/标题样式）:')
    cnt = 0
    for p in doc.paragraphs:
        sn = p.style.name
        if 'Heading' in sn or '标题' in sn or sn in ('1','2','3'):
            print(f'    [{sn}] "{p.text[:35]}"')
            cnt += 1
            if cnt >= 12: break
    print(f'  Normal段落前5:')
    ncnt = 0
    for p in doc.paragraphs:
        if p.style.name == 'Normal' and p.text.strip():
            print(f'    [Normal] "{p.text[:35]}"')
            ncnt += 1
            if ncnt >= 5: break
