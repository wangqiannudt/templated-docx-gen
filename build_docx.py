#!/usr/bin/env python3
"""基于模板docx生成正式文档：保留封面/变更记录/节结构，正文用模板样式（标题自动编号）"""
import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def strip_title_num(text):
    """去掉阿拉伯数字编号前缀（1/1.1，试验类Heading自动编号）；保留中文编号（一、/（一），总结报告Heading无自动编号需文字带编号）"""
    text = re.sub(r'^\d+(\.\d+)*\s+', '', text)
    return text.strip()

def is_index_result(style_name, text):
    """判断段落是否是目录/表图目录的旧显示值行（需清理域值、保留域结构）。"""
    s = (style_name or '').lower()
    t = (text or '').strip()
    return (
        s.startswith('toc')
        or 'table of figures' in s
        or t == 'No table of figures entries found.'
    )

def match_caption(text):
    """识别题注：'图N说明'/'表N说明'。N 为阿拉伯数字，后可跟 空格/./、/． 分隔再接非空说明。
    返回 (label, desc) 或 None。要求'图/表'后必有数字，避免误伤'图片说明'等正文。"""
    m = re.match(r'^(图|表)\s*\d+[\s.、．]*\s*(\S.*)', text)
    if not m:
        return None
    return (m.group(1), m.group(2).strip())

def compute_col_widths(header, ncols, narrow_keys, total_w, narrow_w):
    """计算各列宽(dxa)。表头命中 narrow_keys 的列给 narrow_w，其余均分剩余。
    header 不足 ncols 时按空串处理。"""
    col_w = []
    for ci in range(ncols):
        cell_header = header[ci] if ci < len(header) else ''
        col_w.append(narrow_w if any(k in cell_header for k in narrow_keys) else 0)
    wide_n = col_w.count(0)
    if wide_n > 0:
        wide_w = (total_w - sum(col_w)) // wide_n
        col_w = [w if w > 0 else wide_w for w in col_w]
    return col_w

def detect_page_body_width(doc):
    """从模板第一节 sectPr 读版心宽(dxa) = pgSz.w − pgMar(left+right)。
    读不到则回退 8505（原默认）。"""
    try:
        sec = doc.sections[0]
        pw = int(sec.page_width.twips) if sec.page_width else 0
        lm = int(sec.left_margin.twips) if sec.left_margin else 0
        rm = int(sec.right_margin.twips) if sec.right_margin else 0
        if pw - lm - rm > 0:
            return pw - lm - rm
    except Exception:
        pass
    return 8505

def find_style(doc, candidates):
    """在 doc.styles 里按候选名顺序找第一个存在的样式名；找不到返回 None。"""
    existing = set()
    for s in doc.styles:
        try:
            existing.add(s.name)
        except Exception:
            pass
    for name in candidates:
        if name in existing:
            return name
    return None

def detect_change_log_table(doc):
    """找变更记录表：首行某格含 版本/版本号/修订 的第一个表。
    返回 (table_index, ncols) 或 None。"""
    keywords = ('版本号', '版本', '修订')
    for i, t in enumerate(doc.tables):
        if not t.rows:
            continue
        header_text = ''.join(c.text for c in t.rows[0].cells)
        if any(k in header_text for k in keywords):
            ncols = len(t.rows[0].cells)
            return (i, ncols)
    return None

def apply_change_log(table, rows):
    """重建变更记录表数据行（保留表头行）。rows 为 [[c1,c2,...], ...]；
    按表实际列数对齐，多余列留空，行不足则在表头后追加新行。
    rows 为空则不动表（最小惊讶）。"""
    if not rows:
        return
    ncols = len(table.rows[0].cells)
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    for rd in rows:
        new_tr = deepcopy(table.rows[0]._tr)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        for ci, cell in enumerate(new_row.cells):
            if ci >= ncols:
                break
            val = rd[ci] if ci < len(rd) else ''
            # 清掉继承自表头 run 的格式：移除所有 run 后新增纯净 run 写值
            for para in cell.paragraphs:
                for r in list(para.runs):
                    r._r.getparent().remove(r._r)
            cell.paragraphs[0].add_run(val)

def load_manifest(path):
    """读 manifest YAML，返回 dict；path 为 None 返回 {}。"""
    if not path:
        return {}
    import yaml
    with open(path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f) or {}

def resolve_style(manifest, key, doc, candidates):
    """样式解析：manifest 未写 key → find_style 候选兜底；显式写 null/空 → None（禁用）；写值 → 用值。"""
    if key in manifest:
        val = manifest[key]
        return val if val else None   # 显式 null/空 → 禁用
    return find_style(doc, candidates)

def add_runs(paragraph, text):
    clean = text.replace('**', '').replace('`', '')
    paragraph.add_run(clean)

def add_seq_field(paragraph, label, result_text='1'):
    """添加SEQ自动编号域（fldChar标准方式），Word打开时随updateFields自动编号"""
    from docx.oxml import OxmlElement
    def mk(fieldtype=None, instr=None, text=None):
        r = OxmlElement('w:r')
        if fieldtype:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), fieldtype); r.append(fc)
        elif instr:
            it = OxmlElement('w:instrText'); it.text = instr; r.append(it)
        elif text is not None:
            t = OxmlElement('w:t'); t.text = text; r.append(t)
        return r
    _begin = mk('begin')
    _begin.find(qn('w:fldChar')).set(qn('w:dirty'), 'true')
    paragraph._p.append(_begin)
    paragraph._p.append(mk(instr=' SEQ ' + label + ' \\* ARABIC '))
    paragraph._p.append(mk('separate'))
    paragraph._p.append(mk(text=str(result_text)))
    paragraph._p.append(mk('end'))

def clear_frontmatter_field_results(doc):
    """清掉模板目录/图目录/表目录的旧显示值，保留域结构，打开文档更新域后重建。"""
    for p in doc.paragraphs:
        try:
            style_name = p.style.name.lower()
        except Exception:
            style_name = ''
        if p.style.name == 'Heading 1':
            break
        text = p.text
        if not is_index_result(style_name, text):
            continue
        for r in p._p.findall(qn('w:r')):
            if r.find(qn('w:instrText')) is not None:
                continue
            for t in r.findall(qn('w:t')):
                t.text = ''

def parse_md(md_text):
    items = []
    lines = md_text.split('\n')
    start = 0
    for idx, line in enumerate(lines):
        if line.startswith('# '):
            start = idx
            break
    lines = lines[start:]
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        img_m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img_m:
            items.append(('image', (img_m.group(1), img_m.group(2))))
        elif line.startswith('# '):
            items.append(('h1', strip_title_num(line[2:])))
        elif line.startswith('## '):
            items.append(('h2', strip_title_num(line[3:])))
        elif line.startswith('### '):
            items.append(('h3', strip_title_num(line[4:])))
        elif line.startswith('|') and '---' not in line and '文档编号' not in line and '版本号' not in line[:20]:
            tbl = []
            while i < len(lines) and lines[i].startswith('|'):
                if '---' not in lines[i]:
                    cells = [c.strip() for c in lines[i].strip('|').split('|')]
                    tbl.append(cells)
                i += 1
            if tbl:
                items.append(('table', tbl))
            continue
        elif line.startswith('- '):
            items.append(('li', line[2:].strip()))
        elif re.match(r'^\d+\.\s', line):
            items.append(('li', re.sub(r'^\d+\.\s*', '', line).strip()))
        elif line.strip() == '' or line.strip() == '\\newpage':
            pass
        else:
            items.append(('p', line.strip()))
        i += 1
    return items

def build(template, md_path, out, *,
          cover_title, cover_date, change_log_rows=None,
          image_width_cm=13, narrow_keys=None, strip_heading_num=None,
          table_total_width=None, manifest=None, compute_heading_num=None):
    doc = Document(template)
    body = doc.element.body

    # 样式名解析（manifest 优先；显式写 null/空 → 禁用；未写 key → find_style 候选兜底）
    mf = manifest or {}
    # strip_heading_num 未显式传 → 取 manifest 的值（让 build 直接调用/skill 也生效）
    if strip_heading_num is None:
        strip_heading_num = mf.get('strip_heading_num', False)
    # compute_heading_num：strip 模板自动编号 + build 自己算阿拉伯多级编号拼进标题文字
    # （模板自动编号在 Word 里不可靠：H2/H3 不按上级重置。文字编号确定性可靠）
    if compute_heading_num is None:
        compute_heading_num = mf.get('compute_heading_num', False)
    cover_title_style = resolve_style(mf, 'cover_title_style', doc, ['封皮标题','报告标题','文档标题','主标题','标题'])
    cover_date_style = resolve_style(mf, 'cover_date_style', doc, ['封皮单位','编制日期','日期'])
    cap_mf = mf.get('caption_styles') or {}
    cap_table = resolve_style(cap_mf, 'table', doc, ['表题注','表题','表格题注'])
    cap_figure = resolve_style(cap_mf, 'figure', doc, ['图题注','图','图片题注'])
    if 'heading_styles' in mf:
        heading_styles = mf['heading_styles']
    else:
        heading_styles = [
            find_style(doc, ['Heading 1']) or 'Heading 1',
            find_style(doc, ['Heading 2']) or 'Heading 2',
            find_style(doc, ['Heading 3']) or 'Heading 3',
        ]
    if not heading_styles:
        heading_styles = ['Heading 1', 'Heading 2', 'Heading 3']

    # 1. 修改封面主标题
    title_changed = False
    if cover_title_style and cover_title:
        for p in doc.paragraphs:
            if p.style.name == cover_title_style and p.text.strip() and not title_changed:
                if p.runs:
                    p.runs[0].text = cover_title
                    for r in p.runs[1:]:
                        r.text = ''
                title_changed = True
                break
    elif cover_title_style and not cover_title:
        print('[warn] 找到封面标题样式但未传 --cover-title，跳过')
    else:
        print('[warn] 未找到封面标题样式，跳过封面标题填入')

    # 2. 修改封面日期
    if cover_date_style and cover_date:
        for p in doc.paragraphs:
            if p.style.name == cover_date_style and ('年' in p.text or re.search(r'\d{4}', p.text)):
                if p.runs:
                    p.runs[0].text = cover_date
                    for r in p.runs[1:]:
                        r.text = ''
                break
    elif cover_date_style and not cover_date:
        print('[warn] 找到封面日期样式但未传 --cover-date，跳过')
    else:
        print('[warn] 未找到封面日期样式，跳过封面日期填入')

    # 3. 变更记录表：清空数据行，重建
    _cl = detect_change_log_table(doc)
    if _cl:
        _idx, _ncols = _cl
        apply_change_log(doc.tables[_idx], change_log_rows)

    # 4. 删除原正文区域（第一个Heading1到结尾）
    # 动态找Heading1的styleId（不同模板可能是Heading1/1/11等）
    style_id_to_name = {}
    h1_sid = None
    for s in doc.styles:
        try:
            style_id_to_name[s.style_id] = s.name
            if s.name == 'Heading 1':
                h1_sid = s.style_id
        except Exception:
            pass
    h1_sid = h1_sid or 'Heading1'
    to_remove = []
    found_h1 = False
    for child in list(body.iterchildren()):
        # 保留 body 末尾的 sectPr（含页面尺寸/页边距），否则后续 add_table 取 sections[-1] 会越界
        if child.tag == qn('w:sectPr'):
            continue
        if child.tag == qn('w:p'):
            pPr = child.find(qn('w:pPr'))
            style_val = ''
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val')) or ''
            style_name = style_id_to_name.get(style_val, '')
            if style_val in (h1_sid, 'Heading1', '1') or style_name in ('Heading 1', '标题 1'):
                found_h1 = True
        if found_h1:
            to_remove.append(child)
    for elem in to_remove:
        body.remove(elem)

    # 5. 解析md，添加正文
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    items = parse_md(md_text)

    # 表格宽度参数（统一从 detect_page_body_width 取，删除裸 8505）
    total_w = table_total_width or detect_page_body_width(doc)
    nk = narrow_keys if narrow_keys is not None else ('序号','数量','时间','阶段','年度','是否','编号','级别','日期','方式','用途','类型')
    narrow_w = 1000

    tbl_style = None
    for s in doc.styles:
        if 'Table Grid' in s.name or '网格' in s.name:
            tbl_style = s.name
            break

    seq_counters = {'图': 0, '表': 0}
    heading_cnt = [0, 0, 0]   # compute 模式的多级标题计数器
    def _heading_prefix(level):
        """compute 模式：返回 '1' / '1.1' / '1.1.1 ' 前缀，并维护/重置下级计数器。"""
        heading_cnt[level-1] += 1
        for i in range(level, 3):
            heading_cnt[i] = 0
        return '.'.join(str(heading_cnt[i]) for i in range(level)) + ' '
    def _add_heading(level, text):
        name = heading_styles[level-1] if level-1 < len(heading_styles) else None
        try:
            p = doc.add_paragraph(style=name) if name else doc.add_paragraph()
        except Exception:
            print(f'[warn] Heading {level} 样式 {name!r} 无效，降级为 Normal')
            p = doc.add_paragraph()
        add_runs(p, text)
        return p

    for typ, content in items:
        if typ == 'h1':
            _add_heading(1, (_heading_prefix(1) if compute_heading_num else '') + content)
        elif typ == 'h2':
            _add_heading(2, (_heading_prefix(2) if compute_heading_num else '') + content)
        elif typ == 'h3':
            _add_heading(3, (_heading_prefix(3) if compute_heading_num else '') + content)
        elif typ == 'p':
            clean = content.replace('**', '').replace('`', '')
            cap = match_caption(clean)
            if cap:
                label, desc = cap
                seq_counters[label] += 1
                p = doc.add_paragraph()
                cap_name = cap_table if label == '表' else cap_figure
                if cap_name:
                    try: p.style = doc.styles[cap_name]
                    except Exception: pass
                p.add_run(label + ' ')
                add_seq_field(p, label, seq_counters[label])
                p.add_run(' ' + desc)
            else:
                p = doc.add_paragraph(style='Normal')
                add_runs(p, content)
        elif typ == 'image':
            caption, img_path = content
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_run = pic_p.add_run()
            import os
            _full = img_path if os.path.isabs(img_path) else os.path.join(os.path.dirname(md_path), img_path)
            try:
                pic_run.add_picture(_full, width=Cm(image_width_cm))
            except Exception as e:
                print('图片插入失败', _full, e)
            cap_p = doc.add_paragraph()
            if cap_figure:
                try: cap_p.style = doc.styles[cap_figure]
                except Exception: pass
            cm = match_caption(caption)
            if cm:
                label, desc = cm
                seq_counters[label] += 1
                cap_p.add_run(label+' ')
                add_seq_field(cap_p, label, seq_counters[label])
                cap_p.add_run(' '+desc)
            else:
                cap_p.add_run(caption)
        elif typ == 'li':
            try:
                p = doc.add_paragraph(style='List Paragraph')
            except Exception:
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Cm(0.74)
            add_runs(p, content)
        elif typ == 'table':
            rows = content
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            try:
                tbl.style = tbl_style or 'Table Grid'
            except Exception:
                pass
            from docx.oxml import OxmlElement
            from docx.shared import Twips
            tblPr_el = tbl._element.tblPr
            if tblPr_el is None:
                tblPr_el = OxmlElement('w:tblPr')
                tbl._element.insert(0, tblPr_el)
            # 边框
            tblBorders = OxmlElement('w:tblBorders')
            for edge in ('top','left','bottom','right','insideH','insideV'):
                el = OxmlElement('w:' + edge)
                el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4'); el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
                tblBorders.append(el)
            tblPr_el.append(tblBorders)
            # 总宽度 + fixed布局 + tblInd=0
            for tag, attrs in (('w:tblW', [('w:w', str(total_w)),('w:type','dxa')]),
                               ('w:tblLayout', [('w:type','fixed')]),
                               ('w:tblInd', [('w:w','0'),('w:type','dxa')])):
                el = OxmlElement(tag)
                for k,v in attrs: el.set(qn(k), v)
                tblPr_el.append(el)
            # 单元格内边距
            tblMar = OxmlElement('w:tblCellMar')
            for edge, w in (('top','60'),('bottom','60'),('left','108'),('right','108')):
                el = OxmlElement('w:'+edge); el.set(qn('w:w'), w); el.set(qn('w:type'),'dxa'); tblMar.append(el)
            tblPr_el.append(tblMar)
            tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 列宽：短列窄，其余均分宽
            header = rows[0] if rows else []
            col_w = compute_col_widths(header, ncols, nk, total_w, narrow_w)
            # 关键：设置 tblGrid 的 gridCol（fixed布局下真正决定列宽的是 gridCol）
            tbl_el = tbl._element
            # 移除旧 grid
            old_grid = tbl_el.find(qn('w:tblGrid'))
            if old_grid is not None:
                tbl_el.remove(old_grid)
            # 构造新 grid，插到 tblPr 之后
            new_grid = OxmlElement('w:tblGrid')
            for w in col_w:
                gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); new_grid.append(gc)
            tblPr_el.addnext(new_grid)
            for ci, w in enumerate(col_w):
                tbl.columns[ci].width = Twips(w)
            # 填充单元格
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell_text = row[ci] if ci < len(row) else ''
                    cell = tbl.rows[ri].cells[ci]
                    cell.width = Twips(col_w[ci])
                    tcPr = cell._tc.get_or_add_tcPr()
                    vA = OxmlElement('w:vAlign'); vA.set(qn('w:val'), 'center'); tcPr.append(vA)
                    cell.text = cell_text.replace('**','').replace('`','')
                    for para in cell.paragraphs:
                        # 清除首行缩进（Normal样式带2字符缩进，表格内不要）
                        pPr = para._p.get_or_add_pPr()
                        ind = pPr.find(qn('w:ind'))
                        if ind is None:
                            ind = OxmlElement('w:ind')
                            pPr.append(ind)
                        ind.set(qn('w:firstLine'), '0')
                        ind.set(qn('w:firstLineChars'), '0')
                        ind.set(qn('w:left'), '0')
                        ind.set(qn('w:leftChars'), '0')
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if ri == 0:
                    trPr = tbl.rows[ri]._tr.get_or_add_trPr()
                    th = OxmlElement('w:tblHeader'); trPr.append(th)
                    for ci in range(ncols):
                        for para in tbl.rows[0].cells[ci].paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in para.runs:
                                r.bold = True

    # 可选：去掉Heading 1/2/3自动编号
    #  - strip_heading_num：模板自动编号不每章重置（研究总结报告），用 md 文字编号
    #  - compute_heading_num：模板自动编号在 Word 里不可靠（H2/H3 不按上级重置），
    #    build 自己算阿拉伯编号拼进文字，故也要 strip 避免双编号
    if strip_heading_num or compute_heading_num:
        # 按 manifest 的 heading_styles 名字剥离 numPr（配置驱动，兼容英文 Heading / 中文 标题）
        targets = set(heading_styles or [])
        for s in doc.styles:
            try:
                if s.name in targets:
                    pPr = s.element.find(qn('w:pPr'))
                    if pPr is not None:
                        numPr = pPr.find(qn('w:numPr'))
                        if numPr is not None:
                            pPr.remove(numPr)
            except Exception:
                pass
    clear_frontmatter_field_results(doc)
    doc.save(out)
    print(f"已生成: {out}")

def main():
    import argparse
    p = argparse.ArgumentParser(description='基于模板 docx 从 markdown 生成规范归档文档')
    p.add_argument('template')
    p.add_argument('md')
    p.add_argument('out')
    p.add_argument('--manifest', default=None)
    p.add_argument('--cover-title', default=None, help='封面标题；模板无封面样式时可不传')
    p.add_argument('--cover-date', default=None, help='封面日期；模板无封面样式时可不传')
    p.add_argument('--change-log', action='append', default=None,
                   help='每行一次，列用 | 分隔，如 "V1.0|2025.12|编制单位|增加|说明"')
    p.add_argument('--image-width-cm', type=float, default=13)
    p.add_argument('--narrow-keys', default=None, help='逗号分隔，如 序号,数量,时间')
    p.add_argument('--strip-heading-num', action='store_true', default=None)
    p.add_argument('--compute-heading-num', action='store_true', default=None,
                   help='strip 模板自动编号 + build 自己算阿拉伯多级编号(1/1.1/1.1.1)拼进标题')
    p.add_argument('--table-total-width', type=int, default=None)
    a = p.parse_args()

    manifest = load_manifest(a.manifest)
    change_log_rows = [row.split('|') for row in a.change_log] if a.change_log else None
    narrow_keys = a.narrow_keys.split(',') if a.narrow_keys else None
    # strip/compute：CLI 传 True 则用 CLI；未传(None) → build 内部取 manifest 值
    build(a.template, a.md, a.out,
          cover_title=a.cover_title, cover_date=a.cover_date,
          change_log_rows=change_log_rows,
          image_width_cm=a.image_width_cm, narrow_keys=narrow_keys,
          strip_heading_num=a.strip_heading_num,
          compute_heading_num=a.compute_heading_num,
          table_total_width=a.table_total_width, manifest=manifest)


if __name__ == '__main__':
    main()
